"""Document processing module for RAG-CLI.

This module handles document loading, chunking, and metadata extraction
for various file formats with semantic boundary preservation.
"""

import os
import re
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple, Union
from datetime import datetime
from dataclasses import dataclass, asdict

# Document parsing imports
import pypdf2
import docx
from bs4 import BeautifulSoup
import markdown

# LangChain for chunking
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.text_splitter import CharacterTextSplitter

from src.core.config import get_config
from src.monitoring.logger import get_logger, get_metrics_logger, log_execution_time


logger = get_logger(__name__)
metrics = get_metrics_logger()


@dataclass
class DocumentChunk:
    """Represents a chunk of a document."""
    content: str
    metadata: Dict[str, Any]
    chunk_index: int
    total_chunks: int
    char_count: int
    token_count: int
    source: str
    doc_id: str
    chunk_id: str


@dataclass
class Document:
    """Represents a complete document."""
    content: str
    source: str
    doc_type: str
    metadata: Dict[str, Any]
    doc_id: str
    timestamp: datetime


class DocumentProcessor:
    """Processes documents for RAG pipeline."""

    def __init__(self):
        """Initialize document processor."""
        config = get_config()
        self.chunk_size = config.document_processing.chunk_size
        self.chunk_overlap = config.document_processing.chunk_overlap
        self.separators = config.document_processing.separators
        self.supported_formats = config.document_processing.supported_formats
        self.add_headers = config.document_processing.add_contextual_headers
        self.metadata_fields = config.document_processing.metadata_fields

        # Initialize text splitter
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=self.separators,
            length_function=self._token_length
        )

        logger.info(
            "Document processor initialized",
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap
        )

    def _token_length(self, text: str) -> int:
        """Estimate token count for text.

        Args:
            text: Input text

        Returns:
            Estimated token count
        """
        # Simple estimation: ~4 characters per token
        return len(text) // 4

    @log_execution_time
    def process_document(
        self,
        file_path: Union[str, Path],
        metadata: Optional[Dict[str, Any]] = None
    ) -> Document:
        """Process a single document from file.

        Args:
            file_path: Path to the document
            metadata: Optional metadata to attach

        Returns:
            Processed document
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")

        # Check if format is supported
        suffix = file_path.suffix.lower()
        if suffix not in self.supported_formats:
            raise ValueError(f"Unsupported format: {suffix}")

        logger.info(f"Processing document", path=str(file_path), format=suffix)

        # Load document content
        content = self._load_document(file_path)

        # Generate document ID
        doc_id = self._generate_doc_id(str(file_path))

        # Extract metadata
        doc_metadata = self._extract_metadata(file_path, content)
        if metadata:
            doc_metadata.update(metadata)

        # Create document
        document = Document(
            content=content,
            source=str(file_path),
            doc_type=suffix[1:],  # Remove the dot
            metadata=doc_metadata,
            doc_id=doc_id,
            timestamp=datetime.now()
        )

        logger.info(
            f"Document processed",
            doc_id=doc_id,
            chars=len(content),
            tokens=self._token_length(content)
        )
        metrics.record_success("document_processing")

        return document

    def _load_document(self, file_path: Path) -> str:
        """Load document content based on file type.

        Args:
            file_path: Path to the document

        Returns:
            Document content as text
        """
        suffix = file_path.suffix.lower()

        try:
            if suffix in ['.txt', '.md']:
                return self._load_text_file(file_path)
            elif suffix == '.pdf':
                return self._load_pdf(file_path)
            elif suffix == '.docx':
                return self._load_docx(file_path)
            elif suffix in ['.html', '.htm']:
                return self._load_html(file_path)
            else:
                # Try to load as text
                return self._load_text_file(file_path)

        except Exception as e:
            logger.error(f"Failed to load document", path=str(file_path), error=str(e))
            raise

    def _load_text_file(self, file_path: Path) -> str:
        """Load plain text or markdown file.

        Args:
            file_path: Path to the file

        Returns:
            File content
        """
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

        # If markdown, optionally convert to plain text
        if file_path.suffix.lower() == '.md':
            # Keep markdown as-is for now (maintains structure)
            pass

        return content

    def _load_pdf(self, file_path: Path) -> str:
        """Load PDF file content.

        Args:
            file_path: Path to the PDF

        Returns:
            Extracted text
        """
        text_parts = []

        with open(file_path, 'rb') as f:
            pdf_reader = pypdf2.PdfReader(f)
            num_pages = len(pdf_reader.pages)

            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:
                    text_parts.append(text)

        return '\n\n'.join(text_parts)

    def _load_docx(self, file_path: Path) -> str:
        """Load DOCX file content.

        Args:
            file_path: Path to the DOCX

        Returns:
            Extracted text
        """
        doc = docx.Document(file_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        return '\n\n'.join(text_parts)

    def _load_html(self, file_path: Path) -> str:
        """Load HTML file content.

        Args:
            file_path: Path to the HTML

        Returns:
            Extracted text
        """
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            html_content = f.read()

        soup = BeautifulSoup(html_content, 'html.parser')

        # Remove script and style elements
        for element in soup(['script', 'style']):
            element.decompose()

        # Get text content
        text = soup.get_text()

        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)

        return text

    def _extract_metadata(self, file_path: Path, content: str) -> Dict[str, Any]:
        """Extract metadata from document.

        Args:
            file_path: Path to the document
            content: Document content

        Returns:
            Metadata dictionary
        """
        metadata = {
            'filename': file_path.name,
            'file_path': str(file_path),
            'file_size': file_path.stat().st_size,
            'modified_time': datetime.fromtimestamp(file_path.stat().st_mtime),
            'char_count': len(content),
            'token_count': self._token_length(content),
            'line_count': content.count('\n') + 1
        }

        # Extract title if possible
        title = self._extract_title(file_path, content)
        if title:
            metadata['title'] = title

        # Extract sections for markdown
        if file_path.suffix.lower() == '.md':
            sections = self._extract_markdown_sections(content)
            metadata['sections'] = sections

        return metadata

    def _extract_title(self, file_path: Path, content: str) -> Optional[str]:
        """Extract title from document.

        Args:
            file_path: Path to the document
            content: Document content

        Returns:
            Title if found
        """
        # For markdown, look for # heading
        if file_path.suffix.lower() == '.md':
            match = re.match(r'^#\s+(.+)$', content, re.MULTILINE)
            if match:
                return match.group(1).strip()

        # For HTML, look for title tag
        if file_path.suffix.lower() in ['.html', '.htm']:
            soup = BeautifulSoup(content, 'html.parser')
            title_tag = soup.find('title')
            if title_tag:
                return title_tag.get_text().strip()

        # Default to filename without extension
        return file_path.stem

    def _extract_markdown_sections(self, content: str) -> List[str]:
        """Extract section headings from markdown.

        Args:
            content: Markdown content

        Returns:
            List of section headings
        """
        sections = []
        for line in content.split('\n'):
            if line.startswith('#'):
                # Remove # symbols and strip
                section = re.sub(r'^#+\s*', '', line).strip()
                if section:
                    sections.append(section)
        return sections

    def _generate_doc_id(self, source: str) -> str:
        """Generate unique document ID.

        Args:
            source: Document source path

        Returns:
            Document ID
        """
        # Use hash of source path and timestamp
        hash_input = f"{source}_{datetime.now().isoformat()}"
        return hashlib.md5(hash_input.encode()).hexdigest()[:16]

    @log_execution_time
    def chunk_document(
        self,
        document: Document,
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None
    ) -> List[DocumentChunk]:
        """Split document into chunks.

        Args:
            document: Document to chunk
            chunk_size: Optional override for chunk size
            chunk_overlap: Optional override for overlap

        Returns:
            List of document chunks
        """
        # Use provided sizes or defaults
        if chunk_size or chunk_overlap:
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size or self.chunk_size,
                chunk_overlap=chunk_overlap or self.chunk_overlap,
                separators=self.separators,
                length_function=self._token_length
            )
        else:
            splitter = self.splitter

        # Split the content
        chunks_text = splitter.split_text(document.content)
        total_chunks = len(chunks_text)

        logger.info(f"Chunking document", doc_id=document.doc_id, chunks=total_chunks)

        # Create chunk objects
        chunks = []
        for i, chunk_text in enumerate(chunks_text):
            # Add contextual header if enabled
            if self.add_headers:
                chunk_text = self._add_contextual_header(document, chunk_text, i)

            # Generate chunk ID
            chunk_id = f"{document.doc_id}_chunk_{i:04d}"

            # Create chunk metadata
            chunk_metadata = {
                **document.metadata,
                'doc_id': document.doc_id,
                'chunk_index': i,
                'total_chunks': total_chunks,
                'source': document.source,
                'doc_type': document.doc_type
            }

            # Create chunk object
            chunk = DocumentChunk(
                content=chunk_text,
                metadata=chunk_metadata,
                chunk_index=i,
                total_chunks=total_chunks,
                char_count=len(chunk_text),
                token_count=self._token_length(chunk_text),
                source=document.source,
                doc_id=document.doc_id,
                chunk_id=chunk_id
            )

            chunks.append(chunk)

        logger.info(
            f"Document chunked",
            doc_id=document.doc_id,
            chunks=len(chunks),
            avg_chars=sum(c.char_count for c in chunks) / len(chunks) if chunks else 0
        )
        metrics.record_count("chunks_created", len(chunks))

        return chunks

    def _add_contextual_header(
        self,
        document: Document,
        chunk_text: str,
        chunk_index: int
    ) -> str:
        """Add contextual header to chunk.

        Args:
            document: Source document
            chunk_text: Chunk content
            chunk_index: Index of chunk

        Returns:
            Chunk with header
        """
        header_parts = []

        # Add document title if available
        if 'title' in document.metadata:
            header_parts.append(f"Document: {document.metadata['title']}")

        # Add source file
        header_parts.append(f"Source: {Path(document.source).name}")

        # Add chunk info
        # header_parts.append(f"Part {chunk_index + 1}")

        if header_parts:
            header = ' | '.join(header_parts)
            return f"[{header}]\n\n{chunk_text}"

        return chunk_text

    @log_execution_time
    def process_directory(
        self,
        directory_path: Union[str, Path],
        recursive: bool = True,
        file_pattern: Optional[str] = None
    ) -> List[Document]:
        """Process all documents in a directory.

        Args:
            directory_path: Path to directory
            recursive: Whether to process subdirectories
            file_pattern: Optional glob pattern for files

        Returns:
            List of processed documents
        """
        directory = Path(directory_path)

        if not directory.is_dir():
            raise NotADirectoryError(f"Not a directory: {directory}")

        # Find files to process
        if recursive:
            if file_pattern:
                files = list(directory.rglob(file_pattern))
            else:
                files = []
                for ext in self.supported_formats:
                    files.extend(directory.rglob(f"*{ext}"))
        else:
            if file_pattern:
                files = list(directory.glob(file_pattern))
            else:
                files = []
                for ext in self.supported_formats:
                    files.extend(directory.glob(f"*{ext}"))

        logger.info(f"Processing directory", path=str(directory), files=len(files))

        # Process each file
        documents = []
        errors = []

        for file_path in files:
            try:
                doc = self.process_document(file_path)
                documents.append(doc)
            except Exception as e:
                logger.error(f"Failed to process file", path=str(file_path), error=str(e))
                errors.append((str(file_path), str(e)))

        # Log summary
        logger.info(
            f"Directory processing complete",
            processed=len(documents),
            errors=len(errors)
        )

        if errors:
            logger.warning(f"Failed files", count=len(errors), files=[e[0] for e in errors])

        return documents

    def process_and_chunk_directory(
        self,
        directory_path: Union[str, Path],
        recursive: bool = True,
        file_pattern: Optional[str] = None
    ) -> Tuple[List[Document], List[DocumentChunk]]:
        """Process and chunk all documents in a directory.

        Args:
            directory_path: Path to directory
            recursive: Whether to process subdirectories
            file_pattern: Optional glob pattern

        Returns:
            Tuple of (documents, chunks)
        """
        # Process documents
        documents = self.process_directory(directory_path, recursive, file_pattern)

        # Chunk all documents
        all_chunks = []
        for doc in documents:
            chunks = self.chunk_document(doc)
            all_chunks.extend(chunks)

        logger.info(
            f"Processed and chunked directory",
            documents=len(documents),
            chunks=len(all_chunks)
        )

        return documents, all_chunks


# Singleton instance
_processor: Optional[DocumentProcessor] = None


def get_document_processor() -> DocumentProcessor:
    """Get or create the global document processor.

    Returns:
        Document processor instance
    """
    global _processor
    if _processor is None:
        _processor = DocumentProcessor()
    return _processor


if __name__ == "__main__":
    # Test document processing
    print("Testing Document Processor...")

    processor = get_document_processor()

    # Create a test markdown file
    test_file = Path("test_document.md")
    test_content = """# RAG System Documentation

## Introduction

This is a test document for the RAG system. It contains multiple sections
and paragraphs to test the chunking functionality.

## Features

The RAG system has several key features:

1. Document processing and chunking
2. Embedding generation
3. Vector storage and retrieval
4. Response generation with Claude

### Document Processing

Documents are split into chunks while preserving semantic boundaries.
This ensures that related information stays together.

### Embedding Generation

We use sentence-transformers to generate embeddings for each chunk.
These embeddings capture the semantic meaning of the text.

## Conclusion

The RAG system provides an efficient way to search and retrieve
relevant information from a large corpus of documents.
"""

    # Write test file
    with open(test_file, 'w') as f:
        f.write(test_content)

    try:
        # Process the document
        print("\nProcessing document...")
        doc = processor.process_document(test_file)
        print(f"Document ID: {doc.doc_id}")
        print(f"Content length: {len(doc.content)} chars")
        print(f"Metadata: {doc.metadata}")

        # Chunk the document
        print("\nChunking document...")
        chunks = processor.chunk_document(doc)
        print(f"Created {len(chunks)} chunks")

        for i, chunk in enumerate(chunks):
            print(f"\nChunk {i + 1}:")
            print(f"  ID: {chunk.chunk_id}")
            print(f"  Chars: {chunk.char_count}, Tokens: {chunk.token_count}")
            print(f"  Content preview: {chunk.content[:100]}...")

    finally:
        # Clean up test file
        if test_file.exists():
            test_file.unlink()

    print("\nDocument processor tests completed!")